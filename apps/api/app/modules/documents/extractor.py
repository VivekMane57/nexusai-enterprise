from dataclasses import dataclass
from pathlib import Path

from docx import Document as DocxDocument
from pypdf import PdfReader

from app.modules.documents.constants import DocumentType


@dataclass(frozen=True)
class ExtractedPage:
    page_number: int
    text: str


@dataclass(frozen=True)
class ExtractedDocument:
    text: str
    page_count: int | None
    pages: list[ExtractedPage]


class DocumentExtractionError(Exception):
    """Raised when document text extraction fails."""


class DocumentTextExtractor:
    """
    Extract readable text from PDF, DOCX and TXT documents.

    Responsibilities:
    - validate that the file exists
    - invoke the correct format-specific extractor
    - normalize extracted text
    - preserve page metadata where available
    """

    def extract(
        self,
        *,
        file_path: str,
        file_type: DocumentType,
    ) -> ExtractedDocument:
        path = Path(file_path)

        if not path.exists():
            raise DocumentExtractionError(
                f"Document file does not exist: {path}"
            )

        if not path.is_file():
            raise DocumentExtractionError(
                f"Document path is not a file: {path}"
            )

        try:
            if file_type == DocumentType.PDF:
                return self._extract_pdf(path)

            if file_type == DocumentType.DOCX:
                return self._extract_docx(path)

            if file_type == DocumentType.TXT:
                return self._extract_txt(path)

        except DocumentExtractionError:
            raise
        except Exception as exc:
            raise DocumentExtractionError(
                f"Failed to extract text from {path.name}: {exc}"
            ) from exc

        raise DocumentExtractionError(
            f"Unsupported document type: {file_type}"
        )

    def _extract_pdf(
        self,
        path: Path,
    ) -> ExtractedDocument:
        reader = PdfReader(str(path))
        extracted_pages: list[ExtractedPage] = []

        for page_index, page in enumerate(
            reader.pages,
            start=1,
        ):
            raw_text = page.extract_text() or ""
            cleaned_text = self._clean_text(raw_text)

            extracted_pages.append(
                ExtractedPage(
                    page_number=page_index,
                    text=cleaned_text,
                )
            )

        combined_text = "\n\n".join(
            page.text
            for page in extracted_pages
            if page.text
        )

        return ExtractedDocument(
            text=combined_text,
            page_count=len(reader.pages),
            pages=extracted_pages,
        )

    def _extract_docx(
        self,
        path: Path,
    ) -> ExtractedDocument:
        document = DocxDocument(str(path))
        blocks: list[str] = []

        for paragraph in document.paragraphs:
            cleaned_text = self._clean_text(
                paragraph.text
            )

            if cleaned_text:
                blocks.append(cleaned_text)

        for table in document.tables:
            for row in table.rows:
                cell_values = [
                    self._clean_text(cell.text)
                    for cell in row.cells
                ]

                non_empty_values = [
                    value
                    for value in cell_values
                    if value
                ]

                if non_empty_values:
                    blocks.append(
                        " | ".join(non_empty_values)
                    )

        combined_text = "\n\n".join(blocks)

        return ExtractedDocument(
            text=combined_text,
            page_count=None,
            pages=[],
        )

    def _extract_txt(
        self,
        path: Path,
    ) -> ExtractedDocument:
        try:
            raw_text = path.read_text(
                encoding="utf-8",
            )
        except UnicodeDecodeError:
            raw_text = path.read_text(
                encoding="latin-1",
            )

        cleaned_text = self._clean_text(raw_text)

        return ExtractedDocument(
            text=cleaned_text,
            page_count=None,
            pages=[],
        )

    @staticmethod
    def _clean_text(text: str) -> str:
        normalized_lines: list[str] = []

        for raw_line in text.splitlines():
            normalized_line = " ".join(
                raw_line.split()
            ).strip()

            if normalized_line:
                normalized_lines.append(
                    normalized_line
                )

        return "\n".join(
            normalized_lines
        ).strip()